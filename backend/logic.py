import os
import json
import re
from difflib import SequenceMatcher
from typing import Dict, List
from langchain_community.document_loaders import PyPDFLoader
from langchain_groq import ChatGroq
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.documents import Document as LCDocument
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader
from sqlalchemy.orm import Session
try:
    from .models import JobDescription, MatchResult, Resume
except ImportError:
    from models import JobDescription, MatchResult, Resume

# Initialize LLM and Embeddings
llm = ChatGroq(model_name="llama-3.3-70b-versatile", temperature=0.0)
embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

VECTOR_STORE_DIR = "vector_stores"
os.makedirs(VECTOR_STORE_DIR, exist_ok=True)

ACTION_VERBS = {
    "built", "designed", "developed", "led", "created", "launched", "scaled", "optimized",
    "improved", "automated", "implemented", "managed", "delivered", "reduced", "increased",
    "analyzed", "architected", "deployed", "mentored", "owned", "drove", "streamlined",
}

LEADERSHIP_TERMS = {
    "led", "managed", "mentored", "owned", "directed", "coached", "hired", "trained",
    "stakeholders", "cross-functional", "team", "roadmap", "strategy",
}

KNOWN_SKILLS = {
    "python", "java", "javascript", "typescript", "sql", "nosql", "postgresql", "mysql",
    "mongodb", "redis", "react", "nextjs", "nodejs", "fastapi", "django", "flask", "spring",
    "aws", "gcp", "azure", "docker", "kubernetes", "terraform", "ci/cd", "git", "linux",
    "pandas", "numpy", "scikit-learn", "machine learning", "deep learning", "nlp", "pytorch",
    "tensorflow", "tableau", "power bi", "excel", "airflow", "spark", "hadoop", "snowflake",
    "databricks", "graphql", "rest", "microservices", "prompt engineering", "llm", "langchain",
}

SKILL_GROUPS = {
    "engineering": {"python", "java", "javascript", "typescript", "react", "nodejs", "graphql", "rest"},
    "data": {"sql", "pandas", "numpy", "scikit-learn", "tableau", "power bi", "spark", "airflow"},
    "ml-ai": {"machine learning", "deep learning", "nlp", "pytorch", "tensorflow", "llm", "langchain"},
    "cloud-devops": {"aws", "gcp", "azure", "docker", "kubernetes", "terraform", "ci/cd", "linux"},
}


def _safe_text(value) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _safe_list(value) -> list:
    if isinstance(value, list):
        return value
    return []


def _infer_summary_from_text(text: str, max_sentences: int = 2) -> str:
    clean = re.sub(r"\s+", " ", text or "").strip()
    if not clean:
        return ""
    parts = re.split(r"(?<=[.!?])\s+", clean)
    if not parts:
        return ""
    return " ".join(parts[:max_sentences]).strip()[:420]


def _to_plain_text(text: str) -> str:
    # Remove common markdown markers so UI doesn't show raw ** or headings.
    output = _safe_text(text)
    output = re.sub(r"\*\*(.*?)\*\*", r"\1", output)
    output = re.sub(r"^#{1,6}\s*", "", output, flags=re.MULTILINE)
    output = re.sub(r"^[\-\*\+]\s+", "", output, flags=re.MULTILINE)
    return output

def _extract_pdf_text_with_fallback(file_path: str):
    """
    Extract text from PDF using PyPDFLoader first, then pypdf fallback.
    Returns (pages, full_text).
    """
    pages = []
    text = ""

    # Primary path: LangChain loader (keeps compatibility with downstream vector indexing)
    try:
        loader = PyPDFLoader(file_path)
        pages = loader.load()
        text = "\n".join((p.page_content or "").strip() for p in pages).strip()
    except Exception as e:
        print(f"PyPDFLoader failed: {e}")

    # Fallback path: raw pypdf extraction
    if not text:
        try:
            reader = PdfReader(file_path)
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception:
                    pass

            fallback_chunks = []
            for i, page in enumerate(reader.pages, start=1):
                page_text = (page.extract_text() or "").strip()
                if page_text:
                    fallback_chunks.append(page_text)
                    pages.append(LCDocument(page_content=page_text, metadata={"source": file_path, "page": i}))

            text = "\n".join(fallback_chunks).strip()
        except Exception as e:
            print(f"pypdf fallback failed: {e}")

    return pages, text

def _fallback_structured_data(text: str):
    """
    Minimal deterministic fallback when LLM JSON parsing fails.
    """
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    email_match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    phone_match = re.search(r"(\+?\d[\d\-\s()]{7,}\d)", text)

    cgpa_match = re.search(r"\b(?:cgpa|gpa)\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?(?:/[0-9]+(?:\.[0-9]+)?)?)", text, flags=re.IGNORECASE)
    return {
        "name": lines[0] if lines else "",
        "email": email_match.group(0) if email_match else "",
        "phone": phone_match.group(0) if phone_match else "",
        "summary": _infer_summary_from_text(text),
        "skills": [],
        "experience": [],
        "education": [{"cgpa": cgpa_match.group(1)}] if cgpa_match else [],
    }


def _extract_docx_text(file_path: str):
    doc = Document(file_path)
    lines = []
    pages = []
    for p in doc.paragraphs:
        value = (p.text or "").strip()
        if value:
            lines.append(value)
    text = "\n".join(lines).strip()
    if text:
        pages = [LCDocument(page_content=text, metadata={"source": file_path, "page": 1})]
    return pages, text


def normalize_skill(skill: str) -> str:
    return re.sub(r"\s+", " ", (skill or "").strip()).lower()


def calculate_semantic_score(text_a: str, text_b: str) -> float:
    if not text_a or not text_b:
        return 0.0
    return round(SequenceMatcher(None, text_a.lower(), text_b.lower()).ratio(), 4)


def _collect_resume_bullets(resume_data: dict, raw_text: str) -> List[str]:
    bullets: List[str] = []
    for exp in _safe_list(resume_data.get("experience")):
        if not isinstance(exp, dict):
            continue
        for key in ("description", "highlights", "impact", "achievements", "responsibilities"):
            value = exp.get(key)
            if isinstance(value, list):
                for item in value:
                    text = _safe_text(item)
                    if text:
                        bullets.append(text)
            else:
                text = _safe_text(value)
                if text:
                    for line in text.splitlines():
                        line = line.strip().lstrip("-*• ").strip()
                        if line:
                            bullets.append(line)

    if bullets:
        return bullets

    for line in (raw_text or "").splitlines():
        line = line.strip()
        if re.match(r"^[-*•]\s+", line):
            cleaned = re.sub(r"^[-*•]\s+", "", line).strip()
            if cleaned:
                bullets.append(cleaned)

    return bullets


def _extract_jd_skills(job_description: str) -> List[str]:
    text = (job_description or "").lower()
    if not text:
        return []

    found = set()

    # Multi-word first to avoid missing compound skills.
    for skill in sorted(KNOWN_SKILLS, key=len, reverse=True):
        if " " in skill and skill in text:
            found.add(skill)

    tokens = re.findall(r"[a-zA-Z][a-zA-Z0-9/+#\.-]{1,30}", text)
    for token in tokens:
        normalized = normalize_skill(token)
        if normalized in KNOWN_SKILLS:
            found.add(normalized)

    return sorted(found)


def _score_resume_dna(skills: List[str], bullets: List[str], raw_text: str) -> Dict[str, int]:
    raw = raw_text or ""
    lower_raw = raw.lower()
    word_count = len(re.findall(r"\b\w+\b", raw))

    quantified = 0
    action_verb_hits = 0
    avg_words = 0

    if bullets:
        quantified = sum(1 for b in bullets if re.search(r"(\d+%|\$\d+|\d+\+?)", b))
        action_verb_hits = sum(1 for b in bullets if any(re.search(rf"\b{verb}\b", b.lower()) for verb in ACTION_VERBS))
        avg_words = int(sum(len(re.findall(r"\b\w+\b", b)) for b in bullets) / len(bullets))

    unique_skills = {normalize_skill(s) for s in skills if _safe_text(s)}
    domain_coverage = 0
    for _, group in SKILL_GROUPS.items():
        if unique_skills.intersection(group):
            domain_coverage += 1

    impact = min(100, int((quantified / max(1, len(bullets))) * 70 + min(len(unique_skills), 15) * 2))
    momentum = min(100, int((action_verb_hits / max(1, len(bullets))) * 60 + min(len(bullets), 10) * 4))
    depth = min(100, int(min(len(unique_skills), 18) * 4 + domain_coverage * 8))

    # Best readability window: ~12-28 words per bullet.
    if avg_words == 0:
        clarity = 35
    elif 12 <= avg_words <= 28:
        clarity = 90
    elif 8 <= avg_words < 12 or 28 < avg_words <= 36:
        clarity = 72
    else:
        clarity = 55

    leadership_hits = sum(1 for term in LEADERSHIP_TERMS if re.search(rf"\b{re.escape(term)}\b", lower_raw))
    leadership = min(100, 20 + leadership_hits * 12)

    # Penalize extremely short resumes.
    if word_count < 220:
        impact = max(impact - 12, 5)
        momentum = max(momentum - 8, 5)

    return {
        "impact": impact,
        "momentum": momentum,
        "technical_depth": depth,
        "clarity": clarity,
        "leadership": leadership,
    }


def build_innovation_lab(resume_data: dict, raw_text: str, job_description: str = "") -> dict:
    """
    Builds a differentiated insight pack with actionability:
    - Career DNA scores
    - Signal flags
    - Ghost-gap simulation against optional JD
    - One-click edit moves
    """
    structured = resume_data if isinstance(resume_data, dict) else {}
    skills = [s for s in _safe_list(structured.get("skills")) if _safe_text(s)]
    bullets = _collect_resume_bullets(structured, raw_text)
    dna = _score_resume_dna(skills, bullets, raw_text or "")

    quantified_ratio = 0.0
    if bullets:
        quantified_ratio = sum(1 for b in bullets if re.search(r"(\d+%|\$\d+|\d+\+?)", b)) / len(bullets)

    signal_flags = []
    if quantified_ratio < 0.35:
        signal_flags.append(
            {
                "label": "Low quantified impact",
                "severity": "high",
                "detail": "Less than 35% of bullets have numbers, percentages, or measurable outcomes.",
            }
        )
    if len(skills) < 8:
        signal_flags.append(
            {
                "label": "Narrow skill surface",
                "severity": "medium",
                "detail": "Fewer than 8 explicit skills detected. Expand tools, frameworks, and domain keywords.",
            }
        )
    if dna["leadership"] < 45:
        signal_flags.append(
            {
                "label": "Leadership signal is weak",
                "severity": "medium",
                "detail": "Resume has limited evidence of ownership, mentoring, or cross-functional leadership.",
            }
        )

    jd_skills = _extract_jd_skills(job_description)
    resume_skill_set = {normalize_skill(s) for s in skills}
    matched = [s for s in jd_skills if normalize_skill(s) in resume_skill_set]
    missing = [s for s in jd_skills if normalize_skill(s) not in resume_skill_set]

    current_score = round((len(matched) / len(jd_skills)) * 100, 2) if jd_skills else 0.0
    projected_gains = []
    if jd_skills:
        per_skill_gain = round(100 / len(jd_skills), 2)
        for skill in missing[:5]:
            projected_gains.append(
                {
                    "skill": skill,
                    "projected_score_if_added": min(100.0, round(current_score + per_skill_gain, 2)),
                    "delta": per_skill_gain,
                }
            )

    next_best_moves = [
        {
            "title": "Quantified Bullet Upgrade",
            "instruction": "Rewrite my experience bullets to include measurable impact using %, $, and time-saved metrics.",
            "estimated_gain": max(8, int((0.35 - quantified_ratio) * 40)) if quantified_ratio < 0.35 else 6,
        },
        {
            "title": "Signal Leadership",
            "instruction": "Strengthen leadership evidence in my summary and experience by emphasizing ownership, mentoring, and cross-functional outcomes.",
            "estimated_gain": 7 if dna["leadership"] < 55 else 4,
        },
    ]

    if missing:
        next_best_moves.append(
            {
                "title": "Ghost-Gap Fill",
                "instruction": f"Integrate these JD keywords authentically where relevant: {', '.join(missing[:4])}.",
                "estimated_gain": min(18, 5 + len(missing[:4]) * 3),
            }
        )

    next_best_moves.append(
        {
            "title": "Clarity Compression",
            "instruction": "Compress long bullets into concise STAR-style bullets with one result sentence each.",
            "estimated_gain": 5 if dna["clarity"] < 70 else 3,
        }
    )

    interview_story_arcs = []
    for idx, bullet in enumerate(bullets[:3], start=1):
        interview_story_arcs.append(
            {
                "title": f"Story Arc {idx}",
                "prompt": f"Turn this into STAR format and emphasize trade-offs and impact: {bullet[:180]}",
            }
        )

    signature = []
    normalized_skills = sorted({normalize_skill(s) for s in skills})
    for group_name, group_skills in SKILL_GROUPS.items():
        overlap = sorted(set(normalized_skills).intersection(group_skills))
        if overlap:
            signature.append({"domain": group_name, "skills": overlap[:4]})

    return {
        "career_dna": dna,
        "signal_flags": signal_flags,
        "ghost_gap_simulation": {
            "jd_skills_detected": jd_skills,
            "current_keyword_match_score": current_score,
            "top_missing_skills": missing[:8],
            "projected_gains": projected_gains,
        },
        "next_best_moves": next_best_moves[:4],
        "interview_story_arcs": interview_story_arcs,
        "uniqueness_signature": signature[:4],
    }

def _normalize_structured_data(data: dict):
    """
    Normalize parser output to a stable schema the frontend expects.
    """
    if not isinstance(data, dict):
        data = {}

    normalized_education = []
    for edu in _safe_list(data.get("education", [])):
        if not isinstance(edu, dict):
            continue
        normalized_education.append(
            {
                "degree": _safe_text(edu.get("degree", "")),
                "school": _safe_text(edu.get("school", "")),
                "year": _safe_text(edu.get("year", "")),
                "cgpa": _safe_text(edu.get("cgpa", edu.get("gpa", ""))),
                "location": _safe_text(edu.get("location", "")),
                "details": _safe_text(edu.get("details", "")),
            }
        )

    normalized_experience = []
    for exp in _safe_list(data.get("experience", [])):
        if not isinstance(exp, dict):
            continue
        normalized_experience.append(
            {
                "title": _safe_text(exp.get("title", "")),
                "company": _safe_text(exp.get("company", "")),
                "dates": _safe_text(exp.get("dates", "")),
                "location": _safe_text(exp.get("location", "")),
                "description": _safe_text(exp.get("description", "")),
            }
        )

    normalized_projects = []
    for proj in _safe_list(data.get("projects", [])):
        if not isinstance(proj, dict):
            continue
        normalized_projects.append(
            {
                "title": _safe_text(proj.get("title", "")),
                "dates": _safe_text(proj.get("dates", "")),
                "description": _safe_text(proj.get("description", "")),
            }
        )

    normalized_links = []
    for link in _safe_list(data.get("links", [])):
        if isinstance(link, dict):
            normalized_links.append({
                "title": _safe_text(link.get("title", "")),
                "url": _safe_text(link.get("url", "")),
            })

    normalized = {
        "name": _safe_text(data.get("name", "")),
        "email": _safe_text(data.get("email", "")),
        "phone": _safe_text(data.get("phone", "")),
        "location": _safe_text(data.get("location", "")),
        "summary": _safe_text(data.get("summary", "")),
        "skills": _safe_list(data.get("skills", [])),
        "skill_categories": data.get("skill_categories", {}),
        "links": normalized_links,
        "experience": normalized_experience,
        "education": normalized_education,
        "projects": normalized_projects,
    }

    return normalized


def enrich_structured_data_with_text_hints(data: dict, raw_text: str) -> dict:
    normalized = _normalize_structured_data(data)
    text = raw_text or ""
    cgpa_match = re.search(r"\b(?:cgpa|gpa)\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?(?:/[0-9]+(?:\.[0-9]+)?)?)", text, flags=re.IGNORECASE)
    if not cgpa_match:
        return normalized

    cgpa_value = cgpa_match.group(1)
    education = normalized.get("education", [])
    if education:
        for edu in education:
            if isinstance(edu, dict) and not _safe_text(edu.get("cgpa")):
                edu["cgpa"] = cgpa_value
    else:
        normalized["education"] = [{"degree": "", "school": "", "year": "", "cgpa": cgpa_value, "location": "", "details": ""}]

    return normalized

def parse_resume(file_path: str):
    """
    Parses a PDF resume, extracts structured data, and builds a vector index.
    """
    if file_path.lower().endswith(".pdf"):
        pages, text = _extract_pdf_text_with_fallback(file_path)
    elif file_path.lower().endswith(".docx"):
        pages, text = _extract_docx_text(file_path)
    else:
        raise ValueError("Only PDF and DOCX files are supported for resume parsing.")

    if not text:
        raise ValueError("Could not extract readable text from this file.")
    
    # Build RAG Index
    build_vector_store(pages, file_path)

    # Prompt for structured extraction
    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are an expert resume parser. Extract the following fields from the resume text: name, email, phone, location, summary, skills (list of strings), skill_categories (object grouping skills by category like 'Core Backend', 'Data & Systems', etc.), links (list of objects with 'title' and 'url' fields). For experience, return a list of objects with exactly these keys: 'title', 'company', 'dates', 'location', 'description'. For education, return a list of objects with exactly these keys: 'degree', 'school', 'year', 'cgpa', 'location', 'details'. For projects, return a list of objects with exactly these keys: 'title', 'dates', 'description'. Return ONLY JSON without markdown formatting."),
        ("user", "{text}")
    ])
    
    chain = prompt | llm | JsonOutputParser()
    try:
        structured_data = chain.invoke({"text": text})
        normalized = _normalize_structured_data(structured_data)
        if not normalized.get("summary"):
            normalized["summary"] = _infer_summary_from_text(text)
        return text, normalized
    except Exception as e:
        print(f"Error parsing resume with LLM JSON output: {e}")
        return text, _fallback_structured_data(text)

def build_vector_store(pages, file_path):
    """
    Builds and saves a FAISS vector index for the resume.
    """
    try:
        if not pages:
            print("Skipping vector store build: no extracted pages available.")
            return

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        splits = text_splitter.split_documents(pages)
        if not splits:
            print("Skipping vector store build: no text chunks generated.")
            return
        vectorstore = FAISS.from_documents(documents=splits, embedding=embeddings)
        
        # Save index locally using the filename as ID
        file_name = os.path.basename(file_path)
        index_path = os.path.join(VECTOR_STORE_DIR, file_name)
        vectorstore.save_local(index_path)
        print(f"Vector store saved to {index_path}")
    except Exception as e:
        print(f"Error building vector store: {e}")

def query_resume_context(file_path: str, query: str):
    """
    Retrieves relevant context from the vector store.
    """
    try:
        file_name = os.path.basename(file_path)
        index_path = os.path.join(VECTOR_STORE_DIR, file_name)
        
        if not os.path.exists(index_path):
            return "No context available (Index not found)."
            
        vectorstore = FAISS.load_local(index_path, embeddings, allow_dangerous_deserialization=True)
        retriever = vectorstore.as_retriever(search_kwargs={"k": 3})
        docs = retriever.invoke(query)
        return "\n\n".join([doc.page_content for doc in docs])
    except Exception as e:
        print(f"Error querying context: {e}")
        return "Error retrieving context."

def chat_with_resume(history: list, message: str, resume_data: dict):
    """
    Chat with the resume data.
    """
    context = json.dumps(resume_data, indent=2)
    
    # Format history for LangChain
    chat_history = []
    for msg in history:
        role = msg.get('role')
        content = msg.get('content')
        if role == 'user':
            chat_history.append(("human", content))
        elif role == 'ai':
            chat_history.append(("ai", content))

    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are a helpful AI assistant analyzing a resume. Here is the structured resume data: {context}. Always prioritize resume-specific insights first. If data is missing, state that clearly and suggest what to add. Respond in plain text only (no markdown, no bold markers, no heading markers). Be concise and professional."),
        ("placeholder", "{chat_history}"),
        ("human", "{message}")
    ])
    
    chain = prompt | llm
    response = chain.invoke({
        "context": context,
        "chat_history": chat_history,
        "message": message
    })
    return _to_plain_text(response.content)

def edit_resume(resume_data: dict, instruction: str, context: str = ""):
    """
    Edits the resume data based on user instruction and retrieved context.
    """
    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are a professional resume editor. Update the JSON resume data based on the user's instruction. Use the Retrieved Context from the original PDF to add specific details if needed. Do not fabricate any new accomplishments, skills, dates, locations, employers, or metrics. If the instruction cannot be satisfied without inventing information, preserve the existing content and make only conservative, truthful edits. Return ONLY the valid JSON."),
        ("user", "Current Data: {data}\n\nRetrieved Context: {context}\n\nInstruction: {instruction}")
    ])
    
    chain = prompt | llm | JsonOutputParser()
    updated_data = chain.invoke({"data": json.dumps(resume_data), "context": context, "instruction": instruction})
    return updated_data

def analyze_job_match(resume_data: dict, job_description: str):
    """
    Analyzes the match between a resume and a job description.
    """
    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are an expert HR recruiter. Compare the resume against the job description. Return a JSON object with: 'score' (integer 0-100), 'missing_skills' (list of strings), and 'recommendation' (concise string)."),
        ("user", "Resume Provided: {resume}\n\nJob Description: {job_description}")
    ])
    
    chain = prompt | llm | JsonOutputParser()
    try:
        result = chain.invoke({"resume": json.dumps(resume_data), "job_description": job_description})
        return result
    except Exception as e:
        print(f"Error matching job: {e}")
        return {"score": 0, "missing_skills": [], "recommendation": "Error analyzing match."}

def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_page_margins(section, top: float, right: float, bottom: float, left: float) -> None:
    section.top_margin = Inches(top)
    section.right_margin = Inches(right)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)


def _split_lines(value) -> List[str]:
    if isinstance(value, list):
        return [_safe_text(item) for item in value if _safe_text(item)]
    text = _safe_text(value)
    if not text:
        return []
    return [line.strip() for line in text.splitlines() if line.strip()]


def _render_experience_block(container, experiences: List[dict]) -> None:
    for exp in experiences:
        if not isinstance(exp, dict):
            continue
        title = _safe_text(exp.get("title"))
        company = _safe_text(exp.get("company"))
        dates = _safe_text(exp.get("dates"))
        location = _safe_text(exp.get("location"))

        header = container.add_paragraph()
        header.paragraph_format.space_after = Pt(2)
        if title:
            title_run = header.add_run(title)
            title_run.bold = True
        if company:
            company_run = header.add_run(f" | {company}" if title else company)
            company_run.italic = True
        if dates or location:
            meta = " • ".join(filter(None, [dates, location]))
            if meta:
                meta_run = header.add_run(f"\n{meta}")
                meta_run.font.size = Pt(9)

        details = _split_lines(exp.get("description")) or _split_lines(exp.get("highlights"))
        for detail in details:
            bullet = container.add_paragraph(style=None)
            bullet.style = "List Bullet"
            bullet.paragraph_format.space_after = Pt(0)
            bullet.add_run(detail)


def _render_education_block(container, education_items: List[dict]) -> None:
    for edu in education_items:
        if not isinstance(edu, dict):
            continue
        line = " | ".join(filter(None, [_safe_text(edu.get("degree")), _safe_text(edu.get("school"))]))
        if line:
            para = container.add_paragraph()
            run = para.add_run(line)
            run.bold = True
        meta_bits = [
            _safe_text(edu.get("year")),
            _safe_text(edu.get("location")),
            f"CGPA: {_safe_text(edu.get('cgpa'))}" if _safe_text(edu.get("cgpa")) else "",
        ]
        meta_line = " | ".join(bit for bit in meta_bits if bit)
        if meta_line:
            meta_para = container.add_paragraph()
            meta_para.paragraph_format.space_after = Pt(1)
            meta_run = meta_para.add_run(meta_line)
            meta_run.font.size = Pt(9)
        details = _safe_text(edu.get("details"))
        if details:
            details_para = container.add_paragraph()
            details_para.paragraph_format.space_after = Pt(2)
            details_para.add_run(details)


def _generate_classic_resume_docx(data: dict, output_path: str) -> str:
    doc = Document()
    doc.add_heading(data.get("name") or "Name", 0)

    email = data.get("email") or ""
    phone = data.get("phone") or ""
    contact_info = " | ".join(filter(None, [email, phone]))
    if contact_info:
        doc.add_paragraph(contact_info)

    if data.get("summary"):
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(data.get("summary") or "")

    if data.get("experience"):
        doc.add_heading("Experience", level=1)
        _render_experience_block(doc, data.get("experience", []))

    if data.get("education"):
        doc.add_heading("Education", level=1)
        _render_education_block(doc, data.get("education", []))

    if data.get("skills"):
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(data.get("skills") or []))

    doc.save(output_path)
    return output_path


def _generate_premium_resume_docx(data: dict, output_path: str) -> str:
    doc = Document()
    section = doc.sections[0]
    _set_page_margins(section, 0.4, 0.4, 0.4, 0.4)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]
    
    left_cell.width = Inches(2.1)
    right_cell.width = Inches(4.5)
    
    _set_cell_shading(left_cell, "0D1321")

    para = left_cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)

    name_run = para.add_run(_safe_text(data.get("name")) or "Name")
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = RGBColor(255, 255, 255)

    title_text = ""
    experience_items = data.get("experience", []) if isinstance(data.get("experience"), list) else []
    if experience_items and isinstance(experience_items[0], dict):
        title_text = _safe_text(experience_items[0].get("title"))
    
    if title_text:
        title_para = left_cell.add_paragraph()
        title_para.paragraph_format.space_before = Pt(0)
        title_para.paragraph_format.space_after = Pt(8)
        title_run = title_para.add_run(title_text)
        title_run.font.size = Pt(10)
        title_run.font.color.rgb = RGBColor(165, 180, 252)

    def add_sidebar_section(title: str) -> None:
        """Add a section header in the sidebar."""
        sec_para = left_cell.add_paragraph()
        sec_para.paragraph_format.space_before = Pt(10)
        sec_para.paragraph_format.space_after = Pt(6)
        sec_run = sec_para.add_run(title.upper())
        sec_run.bold = True
        sec_run.font.size = Pt(8)
        sec_run.font.color.rgb = RGBColor(100, 200, 200)

    def add_sidebar_text(text: str, bold: bool = False) -> None:
        """Add regular text to sidebar."""
        para = left_cell.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(_safe_text(text))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(226, 232, 240)
        if bold:
            run.bold = True

    add_sidebar_section("CONTACT")
    contact_items = [
        _safe_text(data.get("email")),
        _safe_text(data.get("phone")),
        _safe_text(data.get("location")),
    ]
    for item in contact_items:
        if item:
            add_sidebar_text(item)

    links = data.get("links", [])
    if links and isinstance(links, list):
        add_sidebar_section("LINKS")
        for link_item in links:
            if isinstance(link_item, dict):
                link_text = _safe_text(link_item.get("title")) or _safe_text(link_item.get("name"))
                if link_text:
                    add_sidebar_text(link_text)
            elif isinstance(link_item, str):
                add_sidebar_text(link_item)

    skills = data.get("skills", []) if isinstance(data.get("skills"), list) else []
    skill_categories = data.get("skill_categories", {})
    
    if skills:
        add_sidebar_section("SKILLS")
        if skill_categories and isinstance(skill_categories, dict):
            for category, cat_skills in skill_categories.items():
                cat_para = left_cell.add_paragraph()
                cat_para.paragraph_format.space_before = Pt(4)
                cat_para.paragraph_format.space_after = Pt(1)
                cat_run = cat_para.add_run(f"{category.title()}")
                cat_run.bold = True
                cat_run.font.size = Pt(8)
                cat_run.font.color.rgb = RGBColor(148, 163, 184)
                
                skills_text = ", ".join(cat_skills) if isinstance(cat_skills, list) else str(cat_skills)
                skills_para = left_cell.add_paragraph()
                skills_para.paragraph_format.space_before = Pt(0)
                skills_para.paragraph_format.space_after = Pt(3)
                skills_run = skills_para.add_run(skills_text)
                skills_run.font.size = Pt(8)
                skills_run.font.color.rgb = RGBColor(226, 232, 240)
        else:
            grouped = [", ".join(skills[i : i + 3]) for i in range(0, len(skills), 3)]
            for group in grouped[:5]:
                add_sidebar_text(group)

    def add_main_section_heading(text: str) -> None:
        """Add main content section heading."""
        para = right_cell.add_paragraph()
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(15, 23, 42)
        
        border_para = right_cell.add_paragraph()
        border_para.paragraph_format.space_before = Pt(0)
        border_para.paragraph_format.space_after = Pt(6)
        pPr = border_para._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "0D1321")
        pBdr.append(bottom)
        pPr.append(pBdr)

    if data.get("summary"):
        add_main_section_heading("Professional Summary")
        summary_para = right_cell.add_paragraph()
        summary_para.paragraph_format.space_after = Pt(4)
        summary_para.add_run(_safe_text(data.get("summary")))

    if experience_items:
        add_main_section_heading("Experience")
        _render_experience_block(right_cell, experience_items)

    projects = data.get("projects", [])
    if projects and isinstance(projects, list):
        add_main_section_heading("Projects")
        _render_experience_block(right_cell, projects)

    education_items = data.get("education", []) if isinstance(data.get("education"), list) else []
    if education_items:
        add_main_section_heading("Education")
        _render_education_block(right_cell, education_items)

    doc.add_section(WD_SECTION.CONTINUOUS)
    doc.save(output_path)
    return output_path


def generate_resume_docx(data: dict, output_path: str, template: str = "classic"):
    """
    Generates a DOCX file from the structured resume data.
    """
    selected_template = _safe_text(template).lower() or "classic"
    if selected_template in {"index_html", "html", "premium"}:
        return _generate_premium_resume_docx(data, output_path)
    return _generate_classic_resume_docx(data, output_path)

def get_dashboard_stats(db: Session):
    """
    Calculates statistics for the dashboard.
    """
    def _completeness_score(data: dict) -> int:
        score = 0
        if data.get("email"):
            score += 20
        if data.get("phone"):
            score += 20
        if data.get("skills"):
            score += 30
        if data.get("experience"):
            score += 30
        return score

    total_candidates = db.query(Resume).count()
    active_jobs = db.query(JobDescription).count()
    match_rows = db.query(MatchResult).all()
    all_resumes = db.query(Resume).all()
    if match_rows:
        avg_match_score = round(sum(float(row.match_score) for row in match_rows) / len(match_rows), 2)
    elif all_resumes:
        avg_match_score = round(
            sum(_completeness_score(r.structured_data if isinstance(r.structured_data, dict) else {}) for r in all_resumes) / len(all_resumes),
            2,
        )
    else:
        avg_match_score = 0
    
    # Get recent 5
    recent_resumes = db.query(Resume).order_by(Resume.id.desc()).limit(5).all()
    recent_list = []
    
    for r in recent_resumes:
        data = r.structured_data if isinstance(r.structured_data, dict) else {}
        name = data.get("name", "Unknown Candidate")
        role = data.get("experience", [{}])[0].get("title", "Candidate") if isinstance(data.get("experience"), list) and data.get("experience") and isinstance(data.get("experience")[0], dict) else "Candidate"
        
        latest_match = sorted(r.match_results, key=lambda row: row.id, reverse=True)[0] if r.match_results else None
        if latest_match:
            score = float(latest_match.match_score)
        else:
            score = _completeness_score(data)
        
        recent_list.append({
            "name": name,
            "role": role,
            "score": score,
            "time": "Just now" # timestamps not in model yet, fallback
        })

    # Top Skills Aggregation
    skill_counts = {}
    total_resumes = len(all_resumes)
    
    for r in all_resumes:
        data = r.structured_data if isinstance(r.structured_data, dict) else {}
        skills = data.get("skills", [])
        if isinstance(skills, list):
            for skill in skills:
                if isinstance(skill, str):
                    s = skill.strip().title()
                    skill_counts[s] = skill_counts.get(s, 0) + 1
                
    # Sort and take top 6
    sorted_skills = sorted(skill_counts.items(), key=lambda x: x[1], reverse=True)[:6]
    top_skills = []
    for name, count in sorted_skills:
        percentage = int((count / total_resumes) * 100) if total_resumes > 0 else 0
        top_skills.append({"name": name, "count": count, "percentage": percentage})

    return {
        "total_candidates": total_candidates,
        "resumes_parsed": total_candidates, # Assuming 1:1 for now
        "active_jobs": active_jobs,
        "avg_match_score": avg_match_score,
        "recent_resumes": recent_list,
        "top_skills": top_skills
    }
